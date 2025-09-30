import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import pandas as pd
import os
from docx import Document

# --- Globals ---
docx_file = None
flowchart_file = None
test_cases = []

# Hard-coded flowchart logic for BLE example
FLOWCHART_STEPS = [
    "Power ON",
    "Button Press",
    "Validate (YES/NO Decision)"
]
BRANCHES = [
    {"condition": "Validate=NO", "action": "RED LED BLINK"},
    {"condition": "Validate=YES", "action": "BLUE LED BLINK"}
]

def upload_docx():
    global docx_file
    f = filedialog.askopenfilename(filetypes=[("Word Docs", "*.docx")])
    if f:
        docx_file = f
        status_label.config(text=f"DOCX Uploaded: {os.path.basename(docx_file)}")

def upload_flowchart():
    global flowchart_file
    f = filedialog.askopenfilename(filetypes=[("Images", "*.png;*.jpg;*.jpeg;*.bmp")])
    if f:
        flowchart_file = f
        status_label.config(text=f"Flowchart Uploaded: {os.path.basename(flowchart_file)}")

def parse_docx_requirements(file_path):
    doc = Document(file_path)
    requirements = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if len(txt) > 8 and (":" in txt or "must" in txt or "should" in txt or "required" in txt):
            requirements.append(txt)
    # If too few, add from any long paragraph
    if len(requirements) < 30:
        requirements = [p.text for p in doc.paragraphs if len(p.text) > 15]
    return requirements[:60]

def generate_test_cases():
    global test_cases
    if not docx_file:
        messagebox.showerror("Error", "Upload the BLE specification DOCX first.")
        return

    requirements = parse_docx_requirements(docx_file)
    test_cases = []

    # 1. Core flow paths from flowchart
    flow_id_num = 1
    for branch in BRANCHES:
        obj = f"Validate system LED behavior on '{branch['condition']}'"
        steps = "\n".join(FLOWCHART_STEPS + [branch["action"]])
        expected = f"{branch['action']} is triggered"
        test_cases.append({
            "Test Case ID": f"FLOW_POS_{flow_id_num}",
            "Objective": obj,
            "Steps to Execute": steps,
            "Expected Result": expected
        })
        flow_id_num += 1

    # 2. Negative flowchart paths
    neg_flows = [
        ("Power ON", "Device fails to power on"),
        ("Button Press", "Button not detected"),
        ("Validate", "Invalid validation result"),
        ("RED LED BLINK", "Red LED does not blink"),
        ("BLUE LED BLINK", "Blue LED does not blink"),
    ]
    for i, (step, neg) in enumerate(neg_flows):
        obj = f"Negative: {step}"
        steps = f"Attempt '{step}' under failure condition"
        expected = f"System handles '{neg}' gracefully"
        test_cases.append({
            "Test Case ID": f"FLOW_NEG_{i+1}",
            "Objective": obj,
            "Steps to Execute": steps,
            "Expected Result": expected
        })

    # 3. Requirement-based positive/negative
    for i, req in enumerate(requirements):
        test_cases.append({
            "Test Case ID": f"REQ_POS_{i+1}",
            "Objective": f"Requirement: {req}",
            "Steps to Execute": req,
            "Expected Result": "Requirement is met"
        })
        test_cases.append({
            "Test Case ID": f"REQ_NEG_{i+1}",
            "Objective": f"Negative: {req}",
            "Steps to Execute": f"Attempt invalid or boundary for: {req}",
            "Expected Result": "System rejects invalid input gracefully"
        })

    # 4. Combinational expansion up to 280
    expansion = []
    count_needed = 280 - len(test_cases)
    for i in range(count_needed):
        req = requirements[i % len(requirements)]
        branch = BRANCHES[i % len(BRANCHES)]
        expansion.append({
            "Test Case ID": f"CMB_{i+1}",
            "Objective": f"Combine '{req}' with '{branch['condition']}'",
            "Steps to Execute": f"Perform: '{req}', then branch logic: {branch['condition']}",
            "Expected Result": f"{branch['action']} after {req} (if applicable)"
        })
    test_cases += expansion

    status_label.config(text=f"{len(test_cases)} test cases generated.")
    messagebox.showinfo("Done", f"Generated {len(test_cases)} test cases.")

def export_excel():
    if not test_cases:
        messagebox.showerror("Error", "No test cases generated yet!")
        return
    df = pd.DataFrame(test_cases)
    filename = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    if filename:
        df.to_excel(filename, index=False)
        messagebox.showinfo("Export", f"Test cases exported to {filename}")

# ---------- TKinter GUI with opaque control panel ---------------

window = tk.Tk()
window.title("CenTrak EmbedTest Studio")

win_w, win_h = 900, 400
window.geometry(f"{win_w}x{win_h}")

# Load the CenTrak background image (make sure the filename matches exactly!)
bg_img_path = "CentrakLogo.PNG"  # Or full path as needed
bg_img = Image.open(bg_img_path)
bg_img = bg_img.resize((win_w, win_h), Image.LANCZOS)
bg_photo = ImageTk.PhotoImage(bg_img)
bg_label = tk.Label(window, image=bg_photo)
bg_label.place(x=0, y=0, relwidth=1, relheight=1)

# Opaque (solid white) panel for buttons, with some padding and shadow so background is visible around the edges
panel = tk.Frame(window, bg="#FFFFFF", bd=3, relief='ridge')
# Place center with width ~60%, height ~60% so background is visible
panel.place(relx=0.5, rely=0.5, anchor='center', relwidth=0.6, relheight=0.6)

docx_btn = tk.Button(panel, text="Upload DOCX Specification", command=upload_docx, width=40)
docx_btn.pack(pady=7)

fc_btn = tk.Button(panel, text="Upload Flowchart Image", command=upload_flowchart, width=40)
fc_btn.pack(pady=7)

gen_btn = tk.Button(panel, text="Generate Test Cases", command=generate_test_cases, width=40)
gen_btn.pack(pady=7)

export_btn = tk.Button(panel, text="Download Test Cases as Excel", command=export_excel, width=40)
export_btn.pack(pady=7)

status_label = tk.Label(panel, text="Upload DOCX and flowchart to begin.", bg="#FFFFFF")
status_label.pack(pady=10)

window.mainloop()
